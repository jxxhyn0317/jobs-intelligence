import streamlit as st
import google.generativeai as genai
import json
import re
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 설정 ──────────────────────────────────────────────────────────────────────
GEMINI_API_KEY = "AIzaSyBqf6XTwXdiY-QeShQsMq3aqcr582eB-mM"
MODEL          = "gemini-1.5-pro"  # Pro 업그레이드 — 전략 추론 품질 향상

st.set_page_config(page_title="JD Analyst", page_icon="◎", layout="wide",
                   initial_sidebar_state="collapsed")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;500&family=IBM+Plex+Sans:wght@300;400;500&display=swap');
*, html, body { font-family: 'IBM Plex Sans', sans-serif; }
.stApp { background: linear-gradient(180deg, #000000 0%, #141414 50%, #2a2a2a 100%); color: #d4d0c8; }
.block-container { padding: 0 !important; max-width: 100% !important; }
.hero { padding: 80px 20px 52px; border-bottom: 1px solid #1e1e1e; text-align: center; display: flex; flex-direction: column; align-items: center; }
.hero-eyebrow { font-family: 'IBM Plex Mono', monospace; font-size: 10px; letter-spacing: 0.25em; color: #444; text-transform: uppercase; margin-bottom: 20px; }
.hero-title { font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; font-size: 54px; font-weight: 700; color: #ffffff; letter-spacing: -0.03em; line-height: 1.0; margin-bottom: 18px; text-shadow: 0 0 30px rgba(255,255,255,0.55), 0 0 70px rgba(255,255,255,0.2); }
.hero-desc { font-size: 14px; color: #555; font-weight: 300; line-height: 1.65; text-align: center; }
.form-section { padding: 52px 20px; border-bottom: 1px solid #1e1e1e; display: flex; flex-direction: column; align-items: center; }
.field-label { font-family: 'IBM Plex Mono', monospace; font-size: 10px; letter-spacing: 0.15em; color: #444; text-transform: uppercase; margin-bottom: 8px; display: block; text-align: center; }
.stTextInput { max-width: 584px; margin: 0 auto !important; }
.stTextInput > div > div > input { background: #0d0d0d !important; border: 1px solid #252525 !important; border-radius: 999px !important; color: #d4d0c8 !important; font-family: 'IBM Plex Mono', monospace !important; font-size: 13px !important; padding: 14px 22px !important; transition: border-color 0.2s !important; text-align: center !important; }
.stTextInput > div > div > input:focus { border-color: #ffffff !important; box-shadow: 0 0 20px rgba(255,255,255,0.05) !important; }
.stTextInput > div > div > input::placeholder { color: #2e2e2e !important; text-align: center !important; }
.stButton { max-width: 584px; margin: 0 auto !important; display: flex; justify-content: center; }
.stButton > button { background: #ffffff !important; color: #000000 !important; border: none !important; border-radius: 999px !important; font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif !important; font-size: 12px !important; font-weight: 700 !important; letter-spacing: 0.08em !important; text-transform: uppercase !important; padding: 14px 36px !important; height: auto !important; transition: opacity 0.15s !important; width: 100% !important; box-shadow: 0 0 24px rgba(255,255,255,0.15) !important; }
.stButton > button:hover { opacity: 0.88 !important; }
.stButton > button:disabled { background: #1a1a1a !important; color: #333 !important; box-shadow: none !important; }
.log-box { background: #080808; border: 1px solid #1e1e1e; border-radius: 16px; padding: 20px 28px; font-family: 'IBM Plex Mono', monospace; font-size: 11px; color: #444; line-height: 1.8; margin: 24px auto; max-width: 640px; width: 90%; min-height: 120px; max-height: 320px; overflow-y: auto; }
.log-line { margin: 2px 0; }
.log-ok { color: #7fff7f; }
.log-active { color: #d4d0c8; }
.log-dim { color: #2a2a2a; }
.url-found-box { background: #0d0d0d; border: 1px solid #252525; border-left: 2px solid #ffffff; border-radius: 8px; padding: 12px 20px; margin: 16px auto 0; max-width: 584px; width: 90%; font-family: 'IBM Plex Mono', monospace; font-size: 11px; color: #555; word-break: break-all; text-align: center; }
.result-section { padding: 56px 20px; text-align: center; }
.result-header { font-family: 'IBM Plex Mono', monospace; font-size: 10px; letter-spacing: 0.25em; color: #444; text-transform: uppercase; margin: 0 auto 32px; padding-bottom: 16px; border-bottom: 1px solid #1e1e1e; max-width: 640px; display: flex; justify-content: space-between; align-items: center; }
.stDownloadButton { max-width: 584px; margin: 0 auto !important; }
.stDownloadButton > button { background: transparent !important; border: 1px solid #252525 !important; color: #555 !important; border-radius: 999px !important; font-family: 'IBM Plex Mono', monospace !important; font-size: 10px !important; letter-spacing: 0.1em !important; padding: 10px 24px !important; height: auto !important; width: 100% !important; }
.stDownloadButton > button:hover { border-color: #ffffff !important; color: #ffffff !important; }
div[data-testid="stSidebarNav"] { display: none; }
footer { display: none !important; }
#MainMenu { display: none !important; }
header { display: none !important; }
</style>
""", unsafe_allow_html=True)

# ── Gemini 호출 ────────────────────────────────────────────────────────────────
def ask(prompt: str) -> str:
    try:
        model    = genai.GenerativeModel(MODEL)
        response = model.generate_content(prompt)
        return response.text or ""
    except Exception as e:
        return f"오류: {e}"

def parse_json(text: str):
    """코드블록 제거 후 JSON 추출 — 여러 방법 시도"""
    cleaned = re.sub(r'```(?:json)?\s*', '', text).strip()
    cleaned = re.sub(r'```\s*$', '', cleaned).strip()
    for pattern in [r'\[[\s\S]*\]', r'\{[\s\S]*\}']:
        try:
            m = re.search(pattern, cleaned)
            if m:
                return json.loads(m.group())
        except:
            continue
    return None

# ── Step 0: 채용 URL 탐색 ──────────────────────────────────────────────────────
KNOWN_URLS = {
    "apple": "https://jobs.apple.com", "google": "https://careers.google.com",
    "meta": "https://www.metacareers.com", "amazon": "https://www.amazon.jobs",
    "microsoft": "https://careers.microsoft.com", "netflix": "https://jobs.netflix.com",
    "tesla": "https://www.tesla.com/careers", "openai": "https://openai.com/careers",
    "anthropic": "https://www.anthropic.com/careers", "kakao": "https://careers.kakao.com",
    "naver": "https://recruit.navercorp.com", "samsung": "https://www.samsung.com/us/careers",
    "hyundai": "https://careers.hyundai.com", "lg": "https://careers.lg.com",
    "krafton": "https://krafton.com/careers", "coupang": "https://www.coupang.jobs",
    "line": "https://careers.linecorp.com", "nvidia": "https://www.nvidia.com/en-us/about-nvidia/careers",
}

def find_careers_url(company: str, log) -> str:
    log(f"◎ {company} 채용 페이지 탐색 중...")
    co = company.lower().strip()
    for k, v in KNOWN_URLS.items():
        if k in co:
            log(f"✓ 채용 페이지: {v}", "ok")
            return v
    result = ask(f'"{company}" 회사의 공식 채용 페이지 URL을 JSON으로만: {{"url":"https://..."}}\n다른 텍스트 없이.')
    data = parse_json(result)
    if isinstance(data, dict):
        url = data.get("url","")
        if url and url.startswith("http"):
            log(f"✓ 채용 페이지: {url}", "ok")
            return url
    m = re.search(r'https?://[^\s"\'}\]]+', result)
    if m:
        log(f"✓ URL 발견", "ok")
        return m.group()
    fallback = f"https://www.linkedin.com/jobs/search/?keywords={company.replace(' ','+')}"
    log(f"→ LinkedIn으로 대체", "ok")
    return fallback

# ── Step 1: 공고 목록 수집 (전체) ────────────────────────────────────────────
def collect_all_jobs(company: str, careers_url: str, filter_kw: str, log) -> list:
    log("◎ 전체 채용 공고 목록 수집 중...")
    filter_note = f'"{filter_kw}" 관련 직무만 포함.' if filter_kw else "모든 직무 포함."
    result = ask(f""""{company}" 회사의 채용 공고 목록을 최대한 많이 (최소 15건 이상) 알려줘.
채용 페이지: {careers_url}
{filter_note}
네가 아는 실제 공고를 기반으로 작성해도 돼.
반드시 JSON 배열만 출력 (마크다운 코드블록 없이):
[{{"title":"직함","team":"팀/부서","location":"위치","url":"공고 URL (알면)","date":"날짜"}}]""")
    data = parse_json(result)
    if isinstance(data, list) and len(data) > 0:
        log(f"✓ {len(data)}건 공고 수집 완료", "ok")
        return data
    # 폴백: 직무별로 나눠서 수집
    log("→ 직무별 분할 수집 중...", "active")
    domains = ["Engineering", "Product", "Design", "Data", "Marketing", "Operations", "Sales", "Legal", "HR"]
    all_jobs = []
    for domain in domains[:5]:
        r2 = ask(f""""{company}" {domain} 부서 채용 공고를 JSON 배열로만:
[{{"title":"직함","team":"{domain}","location":"위치"}}]
{filter_note if filter_kw else ""}""")
        d2 = parse_json(r2)
        if isinstance(d2, list):
            all_jobs.extend(d2)
    if all_jobs:
        # 중복 제거
        seen = set()
        unique = [j for j in all_jobs if j.get("title","") not in seen and not seen.add(j.get("title",""))]
        log(f"✓ 총 {len(unique)}건 수집", "ok")
        return unique
    log("→ 기본 목록으로 대체", "dim")
    fk = filter_kw or "Software"
    return [
        {"title": f"Senior {fk} Engineer", "team": "Engineering", "location": "Seoul"},
        {"title": f"{fk} Product Manager", "team": "Product", "location": "Seoul"},
        {"title": f"Staff {fk} Engineer", "team": "Engineering", "location": "Seoul"},
        {"title": f"{fk} Data Scientist", "team": "Data", "location": "Seoul"},
        {"title": f"Principal {fk} Architect", "team": "Engineering", "location": "Seoul"},
    ]

# ── URL에서 JD 텍스트 추출 ────────────────────────────────────────────────────
def fetch_jd_text(url: str) -> str:
    """URL에서 JD 텍스트 추출. JS 렌더링 사이트는 부분적으로만 가능."""
    try:
        import urllib.request
        headers = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"}
        req  = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=10) as resp:
            raw = resp.read().decode("utf-8", errors="ignore")
        # 태그 제거 + 공백 정리
        text = re.sub(r'<script[\s\S]*?</script>', '', raw, flags=re.IGNORECASE)
        text = re.sub(r'<style[\s\S]*?</style>',  '', text, flags=re.IGNORECASE)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text[:8000]  # 앞 8000자
    except:
        return ""

# ── Step 2: 각 공고 내부 심층 분석 ──────────────────────────────────────────
def deep_analyze_jd(company: str, job: dict, idx: int, total: int, log) -> dict:
    title    = job.get("title","")
    team     = job.get("team","")
    location = job.get("location","")
    url      = job.get("url","")

    log(f"  [{idx}/{total}] {title}")

    # URL이 있으면 실제 내용 fetch 시도
    jd_content = ""
    if url and url.startswith("http"):
        jd_content = fetch_jd_text(url)
        if jd_content:
            log(f"    → 실제 JD 내용 {len(jd_content)}자 추출", "ok")

    # 프롬프트: 실제 내용이 있으면 그걸 기반으로, 없으면 모델 지식 활용
    if jd_content:
        content_block = f"""
아래는 실제 채용 공고 페이지에서 추출한 텍스트야. 이 내용을 기반으로 분석해줘:
---
{jd_content}
---"""
    else:
        content_block = f'"{company}"의 "{title}" 채용 공고에 대해 네가 아는 정보를 최대한 상세하게 사용해줘.'

    result = ask(f"""당신은 McKinsey 파트너급 전략 애널리스트입니다. 아래 채용 공고를 분석해서 이 회사가 무엇을 하려는지를 역추적하세요.

회사: {company}
직무: {title} | 팀: {team} | 위치: {location}
{content_block}

분석 지침:
- Job Description, Responsibilities, Qualifications에 나오는 구체적 단어/문구를 직접 인용해서 근거로 삼으세요.
- 단순히 "AI 역량 강화"같은 추상적 판단이 아니라, "이 회사가 구체적으로 무엇을 만들려 하는가"를 추론하세요.
- 요구 기술 스택, 경험 요구사항에서 현재 회사의 기술 수준과 격차를 파악하세요.
- "Preferred Qualifications"는 회사가 가장 원하는 미래 방향을 드러내는 핵심 단서입니다.

반드시 순수 JSON만 출력 (```코드블록 없이):
{{
  "title": "{title}",
  "team": "{team}",
  "location": "{location}",
  "role_summary": "이 직무가 조직 내에서 어떤 역할을 하는지 — 3문장. 단순 업무 설명이 아닌 전략적 맥락 포함.",
  "what_they_will_build": "이 사람이 실제로 만들거나 해결할 것 — 구체적으로 2-3문장",
  "responsibilities_analysis": [
    {{"responsibility": "실제 책임 항목 원문 요약", "what_it_signals": "이것이 시사하는 전략적 의미"}},
    {{"responsibility": "책임 항목 2", "what_it_signals": "전략적 의미 2"}},
    {{"responsibility": "책임 항목 3", "what_it_signals": "전략적 의미 3"}}
  ],
  "qualification_signals": [
    {{"qualification": "자격 요건 원문 요약", "why_it_matters": "이 요건을 요구하는 이유와 전략적 함의"}}
  ],
  "preferred_qual_signals": [
    {{"qualification": "우대 자격 원문", "strategic_intent": "이 우대 자격이 드러내는 회사의 미래 방향"}}
  ],
  "tech_stack": ["실제 언급된 기술1", "기술2", "기술3"],
  "capability_gap": "이 채용이 메우려는 현재 조직의 역량 공백 — 구체적으로",
  "strategic_implication": "이 단 하나의 채용 공고가 시사하는 전략적 의미 — 4-5문장. JD의 구체적 문구를 근거로 인용하면서 분석."
}}""")

    data = parse_json(result)
    if isinstance(data, dict) and "role_summary" in data:
        return data

    return {
        "title": title, "team": team, "location": location,
        "role_summary": f"{company}의 {team} 부서에서 {title}을 채용합니다.",
        "what_they_will_build": f"{team} 관련 핵심 제품 및 시스템 개발.",
        "responsibilities_analysis": [
            {"responsibility": "핵심 제품 개발", "what_it_signals": "주력 서비스 강화"},
            {"responsibility": "크로스펑셔널 협업", "what_it_signals": "조직 통합 추진"},
        ],
        "qualification_signals": [
            {"qualification": "관련 분야 경력 5년+", "why_it_matters": "즉시 전력 투입 가능한 시니어 선호"}
        ],
        "preferred_qual_signals": [
            {"qualification": "글로벌 경험", "strategic_intent": "글로벌 시장 대응 역량 확보"}
        ],
        "tech_stack": [], "capability_gap": f"{team} 분야 전문 역량 보강",
        "strategic_implication": f"{company}의 {team} 역량 강화를 위한 전략적 채용입니다."
    }

# ── Step 3: 전략 통합 분석 ───────────────────────────────────────────────────
def strategic_synthesis(company: str, filter_kw: str, jds: list, log) -> dict:
    log("◎ 전략 통합 분석 및 이니셔티브 예측 중...")

    # JD 분석 결과를 풍부하게 취합
    jd_digest = []
    for j in jds[:25]:
        entry = {
            "title": j.get("title",""),
            "team": j.get("team",""),
            "what_they_will_build": j.get("what_they_will_build",""),
            "capability_gap": j.get("capability_gap",""),
            "strategic_implication": j.get("strategic_implication",""),
            "tech_stack": j.get("tech_stack",[]),
            "preferred_quals": [s.get("strategic_intent","") for s in j.get("preferred_qual_signals",[])],
            "resp_signals": [s.get("what_it_signals","") for s in j.get("responsibilities_analysis",[])]
        }
        jd_digest.append(entry)

    digest_text = json.dumps(jd_digest, ensure_ascii=False)[:9000]

    result = ask(f"""당신은 McKinsey Senior Partner이자 전략 인텔리전스 전문가입니다.
{company}의 채용 공고 {len(jds)}건을 심층 분석한 결과를 바탕으로 전략 보고서를 작성하세요.
{f"분석 범위: {filter_kw} 관련 직무" if filter_kw else ""}

채용 공고 분석 데이터:
{digest_text}

아래 JSON 구조로 분석 결과를 작성하세요. 각 항목은 반드시 구체적이고 실질적이어야 합니다.
단순한 "AI 강화", "디지털 전환" 같은 추상적 표현은 금지입니다.
JD에서 실제로 확인된 단어, 기술, 역할을 근거로 인용하세요.
반드시 순수 JSON만 출력 (```코드블록 없이):

{{
  "one_line_verdict": "{company}가 지금 하려는 것을 한 문장으로 — 구체적으로",

  "executive_summary": "경영진이 읽을 5-6문장 요약. 이 회사가 채용 패턴을 통해 드러내는 전략 방향, 조직 변화, 기술 방향을 구체적으로. JD에서 실제 확인된 내용을 근거로.",

  "strategic_directions": [
    {{
      "direction": "큰 방향성 제목 — 구체적으로 (예: '자체 AI 추론 엔진 내재화', '리테일 미디어 네트워크 구축')",
      "confidence": "高/中/低",
      "evidence_count": 5,
      "evidence_from_jds": "이 방향을 뒷받침하는 JD의 구체적 문구나 역할 3가지",
      "why_now": "왜 지금 이 방향인가 — 시장/경쟁 맥락과 연결",
      "what_success_looks_like": "이 방향이 성공하면 18개월 후 이 회사는 어떤 모습인가"
    }},
    {{
      "direction": "큰 방향성 2",
      "confidence": "高/中/低",
      "evidence_count": 3,
      "evidence_from_jds": "JD 근거",
      "why_now": "맥락",
      "what_success_looks_like": "성공 시 모습"
    }},
    {{
      "direction": "큰 방향성 3",
      "confidence": "中",
      "evidence_count": 2,
      "evidence_from_jds": "JD 근거",
      "why_now": "맥락",
      "what_success_looks_like": "성공 시 모습"
    }}
  ],

  "predicted_initiatives": [
    {{
      "initiative": "구체적 실행 이니셔티브 이름 (예: '자체 LLM 파인튜닝 파이프라인 구축', '리얼타임 개인화 추천 엔진 V2')",
      "likelihood": "High/Medium/Low",
      "timeline": "예상 시작 시점 (예: 6개월 이내, 12개월 이내)",
      "evidence": "이 이니셔티브를 예측하는 JD 근거",
      "what_they_need": "이 이니셔티브에 필요한 인재/기술 (채용 공고에서 확인된 것)",
      "strategic_impact": "이 이니셔티브가 시장에 미칠 영향"
    }},
    {{
      "initiative": "이니셔티브 2",
      "likelihood": "High/Medium/Low",
      "timeline": "예상 시점",
      "evidence": "JD 근거",
      "what_they_need": "필요 역량",
      "strategic_impact": "시장 영향"
    }},
    {{
      "initiative": "이니셔티브 3",
      "likelihood": "Medium",
      "timeline": "예상 시점",
      "evidence": "JD 근거",
      "what_they_need": "필요 역량",
      "strategic_impact": "시장 영향"
    }},
    {{
      "initiative": "이니셔티브 4",
      "likelihood": "Medium",
      "timeline": "예상 시점",
      "evidence": "JD 근거",
      "what_they_need": "필요 역량",
      "strategic_impact": "시장 영향"
    }},
    {{
      "initiative": "이니셔티브 5",
      "likelihood": "Low",
      "timeline": "예상 시점",
      "evidence": "JD 근거",
      "what_they_need": "필요 역량",
      "strategic_impact": "시장 영향"
    }}
  ],

  "capability_map": {{
    "building_now": ["지금 확실히 구축 중인 역량 1 (JD 근거 있음)", "역량 2", "역량 3"],
    "likely_next": ["다음 6-12개월 내 구축할 것으로 예상되는 역량", "역량 2"],
    "notable_absences": ["눈에 띄게 채용하지 않는 역량 — 이 공백의 전략적 의미"]
  }},

  "tech_bets": [
    {{"technology": "기술명", "signal_strength": "Strong/Moderate/Weak", "inference": "이 기술에 베팅하는 이유 추론"}}
  ],

  "competitive_implications": {{
    "for_competitors": "경쟁사는 이 채용 패턴에서 무엇을 경계해야 하는가 — 구체적으로 3-4문장",
    "for_partners": "파트너/벤더는 어떤 기회와 위협을 마주하게 되는가 — 3문장",
    "for_talent_market": "이 채용이 인재 시장에 미치는 영향 — 어떤 스킬셋의 몸값이 오르는가"
  }},

  "key_uncertainties": [
    {{"question": "이 분석에서 가장 불확실한 것", "why_it_matters": "이게 틀리면 분석 전체가 바뀌는 이유"}}
  ]
}}""")

    data = parse_json(result)
    if isinstance(data, dict) and "strategic_directions" in data:
        log("✓ 전략 분석 완료", "ok")
        return data

    log("→ Fallback 분석 적용", "dim")
    teams   = list(set([j.get("team","") for j in jds if j.get("team")]))[:4]
    titles  = [j.get("title","") for j in jds[:5]]
    impls   = [j.get("strategic_implication","") for j in jds[:5] if j.get("strategic_implication")]
    return {
        "one_line_verdict": f"{company}는 {teams[0] if teams else '핵심 기술'} 역량을 내재화하고 제품 경쟁력을 강화하려 한다.",
        "executive_summary": f"{company}은 {len(jds)}건의 채용을 통해 {', '.join(teams[:3])} 분야의 조직 역량을 강화하고 있다. " + (impls[0] if impls else ""),
        "strategic_directions": [
            {"direction": f"{teams[0] if teams else '기술'} 역량 내재화", "confidence": "高",
             "evidence_count": min(len(jds),6),
             "evidence_from_jds": f"{titles[0] if titles else '시니어 엔지니어'} 등 채용 확인",
             "why_now": "경쟁 심화에 따른 핵심 역량 직접 보유 전략",
             "what_success_looks_like": "18개월 내 핵심 기능 외부 의존도 50% 감소"}
        ],
        "predicted_initiatives": [
            {"initiative": f"{teams[0] if teams else '기술'} 플랫폼 자체 구축",
             "likelihood": "High", "timeline": "6-12개월",
             "evidence": f"{titles[0] if titles else '관련 직무'} 채용 공고",
             "what_they_need": f"{teams[0] if teams else '기술'} 전문가",
             "strategic_impact": "시장 경쟁력 강화"}
        ],
        "capability_map": {
            "building_now": teams[:3] if teams else ["기술 역량"],
            "likely_next": ["글로벌 확장", "AI 통합"],
            "notable_absences": ["외부 파트너십 채용 없음 — 내재화 전략 신호"]
        },
        "tech_bets": [{"technology": "AI/ML", "signal_strength": "Moderate", "inference": "다수 JD에서 AI 관련 경험 요구"}],
        "competitive_implications": {
            "for_competitors": f"{company}의 핵심 역량 강화는 경쟁사와의 기술 격차를 확대할 수 있다.",
            "for_partners": "내재화 전략은 기존 파트너십 계약 구조에 압박을 가할 수 있다.",
            "for_talent_market": f"{', '.join(teams[:2]) if teams else '관련 분야'} 경력자 수요 급증 예상."
        },
        "key_uncertainties": [
            {"question": "채용이 신규 사업 빌딩인지 기존 이탈 대체인지",
             "why_it_matters": "이에 따라 전략 방향 해석이 완전히 달라짐"}
        ]
    }

# ── Word 문서 생성 ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_para(doc, text, style=None, bold=False, size=None, color=None,
             space_before=None, space_after=None, align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    if space_after  is not None: p.paragraph_format.space_after  = Pt(space_after)
    if align: p.alignment = align
    run = p.add_run(text)
    if bold:  run.bold = True
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def build_docx(company, filter_kw, careers_url, jds, synthesis, date_str):
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1.2)
    section.top_margin  = section.bottom_margin = Inches(1.0)

    styles = doc.styles
    def S(name, base, size, bold=False, color="1A1A1A", sb=10, sa=5, font="Calibri"):
        try:    s = styles[name]
        except: s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles[base]
        s.font.name  = font
        s.font.size  = Pt(size)
        s.font.bold  = bold
        s.font.color.rgb = RGBColor(*bytes.fromhex(color))
        s.paragraph_format.space_before = Pt(sb)
        s.paragraph_format.space_after  = Pt(sa)
        return s

    S("JD_Cover",   "Normal", 40, bold=True,  color="003366", sb=0,  sa=6,  font="Georgia")
    S("JD_H1",      "Normal", 14, bold=True,  color="003366", sb=18, sa=4,  font="Georgia")
    S("JD_H2",      "Normal", 12, bold=True,  color="1A1A1A", sb=12, sa=3,  font="Georgia")
    S("JD_Lead",    "Normal", 12, bold=False, color="222222", sb=4,  sa=4,  font="Calibri")
    S("JD_Body",    "Normal", 11, bold=False, color="333333", sb=3,  sa=3,  font="Calibri")
    S("JD_Caption", "Normal",  9, bold=False, color="888888", sb=2,  sa=2,  font="Calibri")
    S("JD_Tag",     "Normal",  9, bold=True,  color="FFFFFF", sb=0,  sa=0,  font="Calibri")

    def rule(doc, color="003366", size=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'), str(size))
        bot.set(qn('w:space'),'1');   bot.set(qn('w:color'), color)
        pBdr.append(bot); pPr.append(pBdr)

    def bullet(doc, text, indent=0.3, color="333333"):
        p = doc.add_paragraph(style="JD_Body")
        p.paragraph_format.left_indent   = Inches(indent)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        r = p.add_run(f"▸  {text}")
        r.font.color.rgb = RGBColor(*bytes.fromhex(color))
        return p

    def label_text(doc, label, text, label_color="003366", indent=0.2):
        p = doc.add_paragraph(style="JD_Body")
        p.paragraph_format.left_indent = Inches(indent)
        r1 = p.add_run(f"{label}  ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(*bytes.fromhex(label_color))
        p.add_run(text)
        return p

    # ── 표지 ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph(style="JD_Cover")
    p.add_run(company.upper())

    p2 = doc.add_paragraph(style="JD_Lead")
    r = p2.add_run("Strategic Intelligence Report — Job Description Analysis")
    r.font.color.rgb = RGBColor(85,119,153)
    r.font.bold = True

    doc.add_paragraph()
    for meta in [f"Date: {date_str}", f"Postings Analyzed: {len(jds)}",
                 f"Filter: {filter_kw or 'All Roles'}", f"Source: {careers_url}"]:
        pm = doc.add_paragraph(style="JD_Caption")
        pm.add_run(meta)

    rule(doc)
    doc.add_paragraph()

    # One-line verdict 박스
    verdict = synthesis.get("one_line_verdict","")
    if verdict:
        pv = doc.add_paragraph(style="JD_Lead")
        r1 = pv.add_run("BOTTOM LINE:  ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(0,51,102)
        pv.add_run(verdict)
        pv.paragraph_format.left_indent = Inches(0.2)

    doc.add_page_break()

    # ── PART 1: 경영진 요약 ──────────────────────────────────────────────────
    doc.add_paragraph("PART 1. EXECUTIVE SUMMARY", style="JD_H1")
    rule(doc)
    doc.add_paragraph()
    p_es = doc.add_paragraph(style="JD_Lead")
    p_es.add_run(synthesis.get("executive_summary",""))
    doc.add_paragraph()

    # ── PART 2: 전략 방향성 ──────────────────────────────────────────────────
    doc.add_paragraph("PART 2. STRATEGIC DIRECTIONS", style="JD_H1")
    rule(doc)
    p_intro = doc.add_paragraph(style="JD_Caption")
    p_intro.add_run("채용 공고 분석으로 도출한 이 회사의 큰 전략 방향. 신뢰도는 JD 내 실제 근거 수에 비례.")
    doc.add_paragraph()

    for d in synthesis.get("strategic_directions", []):
        conf = d.get("confidence","中")
        conf_color = {"高":"003366","中":"555555","低":"888888"}.get(conf,"555555")

        p_dh = doc.add_paragraph(style="JD_H2")
        r_c = p_dh.add_run(f"[신뢰도 {conf}]  ")
        r_c.font.color.rgb = RGBColor(*bytes.fromhex(conf_color))
        r_c.font.size = Pt(10)
        p_dh.add_run(d.get("direction",""))

        label_text(doc, "근거 공고:", f"{d.get('evidence_count',0)}건 — {d.get('evidence_from_jds','')}")
        label_text(doc, "왜 지금인가:", d.get("why_now",""))
        label_text(doc, "성공 시 18개월 후:", d.get("what_success_looks_like",""))
        doc.add_paragraph()

    doc.add_page_break()

    # ── PART 3: 예상 실행 이니셔티브 ────────────────────────────────────────
    doc.add_paragraph("PART 3. PREDICTED INITIATIVES", style="JD_H1")
    rule(doc)
    p_i = doc.add_paragraph(style="JD_Caption")
    p_i.add_run("채용 공고에서 역추론한 구체적 실행 이니셔티브. 공식 발표 전 선행 지표로 해석.")
    doc.add_paragraph()

    like_colors = {"High":"C00000","Medium":"7F7F7F","Low":"AAAAAA"}
    for ini in synthesis.get("predicted_initiatives", []):
        lk = ini.get("likelihood","Medium")
        lk_color = like_colors.get(lk,"7F7F7F")

        p_ih = doc.add_paragraph(style="JD_H2")
        r_lk = p_ih.add_run(f"[{lk} Likelihood]  ")
        r_lk.font.color.rgb = RGBColor(*bytes.fromhex(lk_color))
        r_lk.font.size = Pt(10)
        p_ih.add_run(ini.get("initiative",""))

        label_text(doc, "예상 시점:", ini.get("timeline",""))
        label_text(doc, "JD 근거:", ini.get("evidence",""))
        label_text(doc, "필요 역량:", ini.get("what_they_need",""))
        label_text(doc, "시장 영향:", ini.get("strategic_impact",""))
        doc.add_paragraph()

    doc.add_page_break()

    # ── PART 4: 역량 맵 & 기술 베팅 ─────────────────────────────────────────
    doc.add_paragraph("PART 4. CAPABILITY MAP & TECHNOLOGY BETS", style="JD_H1")
    rule(doc)
    doc.add_paragraph()

    cap_map = synthesis.get("capability_map", {})
    p_ch1 = doc.add_paragraph(style="JD_H2")
    p_ch1.add_run("지금 구축 중 (JD에서 확인됨)")
    for item in cap_map.get("building_now", []):
        bullet(doc, item, color="003366")

    doc.add_paragraph()
    p_ch2 = doc.add_paragraph(style="JD_H2")
    p_ch2.add_run("다음 단계 예상")
    for item in cap_map.get("likely_next", []):
        bullet(doc, item, color="555555")

    doc.add_paragraph()
    p_ch3 = doc.add_paragraph(style="JD_H2")
    p_ch3.add_run("주목할 공백 — 채용하지 않는 영역")
    for item in cap_map.get("notable_absences", []):
        bullet(doc, item, color="888888")

    doc.add_paragraph()
    p_tb = doc.add_paragraph(style="JD_H2")
    p_tb.add_run("기술 베팅 신호")
    for tech in synthesis.get("tech_bets", []):
        sig = tech.get("signal_strength","Moderate")
        sig_label = {"Strong":"●●●","Moderate":"●●○","Weak":"●○○"}.get(sig,"●●○")
        p_t = doc.add_paragraph(style="JD_Body")
        p_t.paragraph_format.left_indent = Inches(0.3)
        r_sig = p_t.add_run(f"{sig_label}  {tech.get('technology','')}  ")
        r_sig.bold = True
        p_t.add_run(f"— {tech.get('inference','')}")

    doc.add_page_break()

    # ── PART 5: 공고별 상세 분석 ─────────────────────────────────────────────
    doc.add_paragraph("PART 5. JOB POSTING DEEP ANALYSIS", style="JD_H1")
    rule(doc)
    p_cnt = doc.add_paragraph(style="JD_Caption")
    p_cnt.add_run(f"총 {len(jds)}건 채용 공고 — Description·Responsibilities·Qualifications 기반 심층 분석")
    doc.add_paragraph()

    for i, jd in enumerate(jds):
        p_jt = doc.add_paragraph(style="JD_H2")
        r_i = p_jt.add_run(f"{i+1:02d}.  ")
        r_i.font.color.rgb = RGBColor(0,51,102)
        p_jt.add_run(jd.get("title",""))

        pm = doc.add_paragraph(style="JD_Caption")
        pm.add_run(f"{jd.get('team','')}  ·  {jd.get('location','')}")
        pm.paragraph_format.left_indent = Inches(0.3)
        doc.add_paragraph()

        # 역할 요약
        if jd.get("role_summary"):
            p_rs = doc.add_paragraph(style="JD_Body")
            p_rs.paragraph_format.left_indent = Inches(0.3)
            p_rs.add_run(jd.get("role_summary",""))

        # 무엇을 만드는가
        if jd.get("what_they_will_build"):
            doc.add_paragraph()
            label_text(doc, "이 사람이 실제로 만들 것:", jd.get("what_they_will_build",""), indent=0.3)

        # Responsibilities 분석
        resp_analysis = jd.get("responsibilities_analysis", [])
        if resp_analysis:
            doc.add_paragraph()
            p_rl = doc.add_paragraph(style="JD_H2")
            p_rl.paragraph_format.left_indent = Inches(0.3)
            p_rl.add_run("Responsibilities 분석")
            for ra in resp_analysis:
                p_r = doc.add_paragraph(style="JD_Body")
                p_r.paragraph_format.left_indent = Inches(0.5)
                r1 = p_r.add_run(f"▸ {ra.get('responsibility','')}  ")
                r1.bold = True
                r2 = p_r.add_run(f"→ {ra.get('what_it_signals','')}")
                r2.font.color.rgb = RGBColor(0,51,102)
                r2.font.italic = True

        # Qualifications 분석
        qual_signals = jd.get("qualification_signals", [])
        pref_signals = jd.get("preferred_qual_signals", [])
        if qual_signals or pref_signals:
            doc.add_paragraph()
            p_ql = doc.add_paragraph(style="JD_H2")
            p_ql.paragraph_format.left_indent = Inches(0.3)
            p_ql.add_run("Qualifications 분석")
            for qs in qual_signals:
                p_q = doc.add_paragraph(style="JD_Body")
                p_q.paragraph_format.left_indent = Inches(0.5)
                r1 = p_q.add_run(f"[필수] {qs.get('qualification','')}  ")
                r1.bold = True
                p_q.add_run(f"→ {qs.get('why_it_matters','')}")
            for ps in pref_signals:
                p_p = doc.add_paragraph(style="JD_Body")
                p_p.paragraph_format.left_indent = Inches(0.5)
                r1 = p_p.add_run(f"[우대] {ps.get('qualification','')}  ")
                r1.bold = True
                r1.font.color.rgb = RGBColor(0,51,102)
                p_p.add_run(f"→ {ps.get('strategic_intent','')}")
                p_p.runs[-1].font.color.rgb = RGBColor(0,51,102)

        # 역량 공백 & 전략적 시사점
        if jd.get("capability_gap"):
            doc.add_paragraph()
            label_text(doc, "메우려는 역량 공백:", jd.get("capability_gap",""), indent=0.3)

        if jd.get("strategic_implication"):
            doc.add_paragraph()
            p_si = doc.add_paragraph(style="JD_Body")
            p_si.paragraph_format.left_indent = Inches(0.3)
            p_si.paragraph_format.space_before = Pt(4)
            r_si = p_si.add_run("전략적 시사점:  ")
            r_si.bold = True
            r_si.font.color.rgb = RGBColor(0,51,102)
            r_si.font.size = Pt(11)
            r2 = p_si.add_run(jd.get("strategic_implication",""))
            r2.font.color.rgb = RGBColor(0,51,102)
            r2.font.italic = True

        # Tech stack
        ts = jd.get("tech_stack", [])
        if ts:
            p_ts = doc.add_paragraph(style="JD_Caption")
            p_ts.paragraph_format.left_indent = Inches(0.3)
            p_ts.add_run(f"Tech: {' · '.join(ts)}")

        rule(doc, color="DDDDDD", size=2)
        doc.add_paragraph()

    doc.add_page_break()

    # ── PART 6: 경쟁 시사점 & 불확실성 ──────────────────────────────────────
    doc.add_paragraph("PART 6. COMPETITIVE IMPLICATIONS & UNCERTAINTIES", style="JD_H1")
    rule(doc)
    doc.add_paragraph()

    ci = synthesis.get("competitive_implications", {})
    for label, key in [("경쟁사에 주는 의미", "for_competitors"),
                       ("파트너/벤더에 주는 의미", "for_partners"),
                       ("인재 시장에 주는 의미", "for_talent_market")]:
        p_cl = doc.add_paragraph(style="JD_H2")
        p_cl.add_run(label)
        p_cb = doc.add_paragraph(style="JD_Body")
        p_cb.add_run(ci.get(key,""))
        p_cb.paragraph_format.left_indent = Inches(0.2)
        doc.add_paragraph()

    doc.add_paragraph("핵심 불확실성", style="JD_H2")
    for u in synthesis.get("key_uncertainties", []):
        p_u = doc.add_paragraph(style="JD_Body")
        p_u.paragraph_format.left_indent = Inches(0.2)
        r1 = p_u.add_run(f"Q: {u.get('question','')}  ")
        r1.bold = True
        p_u.add_run(f"→ {u.get('why_it_matters','')}")

    # 푸터
    doc.add_paragraph()
    rule(doc)
    pf = doc.add_paragraph(style="JD_Caption")
    pf.add_run(f"Generated by JD Analyst  ·  {date_str}  ·  {company}  ·  {len(jds)} postings analyzed")
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1.2)
    section.top_margin  = section.bottom_margin = Inches(1.0)

    # 스타일 정의
    styles = doc.styles
    def ensure_style(name, base, size, bold=False, color="1A1A1A", space_b=12, space_a=6):
        try:
            s = styles[name]
        except:
            s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles[base]
        s.font.name  = "Georgia"
        s.font.size  = Pt(size)
        s.font.bold  = bold
        s.font.color.rgb = RGBColor(*bytes.fromhex(color))
        s.paragraph_format.space_before = Pt(space_b)
        s.paragraph_format.space_after  = Pt(space_a)
        return s

    ensure_style("JD_H1",   "Normal",   28, bold=True,  color="003366", space_b=18, space_a=8)
    ensure_style("JD_H2",   "Normal",   16, bold=True,  color="003366", space_b=14, space_a=4)
    ensure_style("JD_H3",   "Normal",   12, bold=True,  color="1A1A1A", space_b=10, space_a=3)
    ensure_style("JD_Body", "Normal",   11, bold=False, color="333333", space_b=3,  space_a=3)
    ensure_style("JD_Meta", "Normal",    9, bold=False, color="888888", space_b=2,  space_a=2)

    # ── 표지 ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    p_title = doc.add_paragraph(style="JD_H1")
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_title.add_run(company.upper())
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p_sub = doc.add_paragraph(style="JD_Body")
    r = p_sub.add_run("Job Description Intelligence Report")
    r.font.size  = Pt(16)
    r.font.color.rgb = RGBColor(85, 119, 153)
    r.font.bold  = True

    doc.add_paragraph()
    meta_items = [
        f"Date of Analysis: {date_str}",
        f"Total Postings Analyzed: {len(jds)}",
        f"Filter Applied: {filter_kw or 'None (All Roles)'}",
        f"Source: {careers_url}",
    ]
    for m in meta_items:
        p = doc.add_paragraph(style="JD_Meta")
        p.add_run(m)

    # 구분선
    doc.add_paragraph()
    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_after = Pt(0)
    p_border = p_rule._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '003366')
    pb.append(bottom); p_border.append(pb)

    doc.add_page_break()

    # ── A. 경영진 요약 ────────────────────────────────────────────────────────
    doc.add_paragraph("A. EXECUTIVE SUMMARY", style="JD_H2")

    p_es = doc.add_paragraph(style="JD_Body")
    p_es.add_run(synthesis.get("executive_summary",""))

    doc.add_paragraph()

    # 핵심 판단
    p_kj = doc.add_paragraph(style="JD_H3")
    p_kj.add_run("핵심 판단")

    conf_labels = {"high": "🔴 HIGH", "medium": "🟡 MEDIUM", "low": "🟢 LOW"}
    for j in synthesis.get("key_judgments", []):
        lvl   = j.get("level","medium")
        label = conf_labels.get(lvl, "●")
        p = doc.add_paragraph(style="JD_Body")
        r1 = p.add_run(f"{label}  ")
        r1.bold = True
        r1.font.size = Pt(10)
        p.add_run(j.get("text",""))
        p.paragraph_format.left_indent = Inches(0.2)

    doc.add_page_break()

    # ── B. 전략 테마 ──────────────────────────────────────────────────────────
    doc.add_paragraph("B. STRATEGIC THEMES", style="JD_H2")

    for theme in synthesis.get("themes", []):
        # 테마 헤더
        p_th = doc.add_paragraph(style="JD_H3")
        conf = theme.get("confidence","中")
        r_num  = p_th.add_run(f"Theme {theme.get('rank','')}  ")
        r_num.font.color.rgb = RGBColor(0,51,102)
        r_name = p_th.add_run(theme.get("name",""))
        r_name.font.color.rgb = RGBColor(0,51,102)

        # 메타
        p_meta = doc.add_paragraph(style="JD_Meta")
        p_meta.add_run(f"신뢰도: {conf}   |   근거 공고: {theme.get('posting_count',0)}건   |   관련 역할: {', '.join(theme.get('roles',[])[:3])}")
        p_meta.paragraph_format.left_indent = Inches(0.2)

        # 근거
        p_ev = doc.add_paragraph(style="JD_Body")
        r_ev = p_ev.add_run("근거:  ")
        r_ev.bold = True
        p_ev.add_run(f'"{theme.get("evidence","")}"')
        p_ev.paragraph_format.left_indent = Inches(0.2)

        # 해석
        p_in = doc.add_paragraph(style="JD_Body")
        r_in = p_in.add_run("전략적 해석:  ")
        r_in.bold = True
        p_in.add_run(theme.get("interpretation",""))
        p_in.paragraph_format.left_indent = Inches(0.2)

        # 대안
        p_alt = doc.add_paragraph(style="JD_Body")
        r_alt = p_alt.add_run("대안 해석:  ")
        r_alt.bold = True
        r_alt.font.color.rgb = RGBColor(136,136,136)
        r2 = p_alt.add_run(theme.get("alternative",""))
        r2.font.color.rgb = RGBColor(136,136,136)
        p_alt.paragraph_format.left_indent = Inches(0.2)

        doc.add_paragraph()

    doc.add_page_break()

    # ── C. 전체 공고 상세 분석 ────────────────────────────────────────────────
    doc.add_paragraph("C. JOB POSTING ANALYSIS", style="JD_H2")
    p_cnt = doc.add_paragraph(style="JD_Body")
    p_cnt.add_run(f"총 {len(jds)}건 채용 공고 심층 분석")
    p_cnt.runs[0].bold = True
    doc.add_paragraph()

    for i, jd in enumerate(jds):
        # 공고 헤더
        p_jt = doc.add_paragraph(style="JD_H3")
        r_idx = p_jt.add_run(f"{i+1:02d}.  ")
        r_idx.font.color.rgb = RGBColor(0,51,102)
        p_jt.add_run(jd.get("title",""))

        # 팀/위치
        p_jm = doc.add_paragraph(style="JD_Meta")
        p_jm.add_run(f"{jd.get('team','')}  ·  {jd.get('location','')}")
        p_jm.paragraph_format.left_indent = Inches(0.3)

        # 요약
        if jd.get("summary"):
            p_js = doc.add_paragraph(style="JD_Body")
            p_js.add_run(jd.get("summary",""))
            p_js.paragraph_format.left_indent = Inches(0.3)

        # 주요 책임
        resps = jd.get("responsibilities", [])
        if resps:
            p_rl = doc.add_paragraph(style="JD_Body")
            r_rl = p_rl.add_run("주요 책임")
            r_rl.bold = True
            p_rl.paragraph_format.left_indent = Inches(0.3)
            for resp in resps:
                p_r = doc.add_paragraph(style="JD_Body")
                p_r.add_run(f"• {resp}")
                p_r.paragraph_format.left_indent = Inches(0.5)
                p_r.paragraph_format.space_after = Pt(1)

        # 자격 요건
        mins = jd.get("min_qualifications", [])
        if mins:
            p_ql = doc.add_paragraph(style="JD_Body")
            r_ql = p_ql.add_run("필수 자격")
            r_ql.bold = True
            p_ql.paragraph_format.left_indent = Inches(0.3)
            for q in mins:
                p_q = doc.add_paragraph(style="JD_Body")
                p_q.add_run(f"• {q}")
                p_q.paragraph_format.left_indent = Inches(0.5)
                p_q.paragraph_format.space_after = Pt(1)

        # 전략적 시사점
        if jd.get("strategic_implication"):
            p_si = doc.add_paragraph(style="JD_Body")
            r_si = p_si.add_run("전략적 시사점:  ")
            r_si.bold = True
            r_si.font.color.rgb = RGBColor(0,51,102)
            r_si2 = p_si.add_run(jd.get("strategic_implication",""))
            r_si2.font.color.rgb = RGBColor(0,51,102)
            p_si.paragraph_format.left_indent = Inches(0.3)

        # 구분선
        p_div = doc.add_paragraph()
        p_div.paragraph_format.space_before = Pt(6)
        p_div.paragraph_format.space_after  = Pt(6)
        pPr = p_div._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'2')
        bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'DDDDDD')
        pBdr.append(bot); pPr.append(pBdr)

    doc.add_page_break()

    # ── D. 기술 투자 신호 ────────────────────────────────────────────────────
    doc.add_paragraph("D. TECHNOLOGY & INVESTMENT SIGNALS", style="JD_H2")
    for tech in synthesis.get("tech_investments", []):
        p_t = doc.add_paragraph(style="JD_Body")
        p_t.add_run(f"• {tech}")
        p_t.paragraph_format.left_indent = Inches(0.2)
    doc.add_paragraph()

    # ── E. 불확실성 ──────────────────────────────────────────────────────────
    doc.add_paragraph("E. UNCERTAINTIES", style="JD_H2")
    for u in synthesis.get("uncertainties", []):
        p_u = doc.add_paragraph(style="JD_Body")
        r_ut = p_u.add_run(f"{u.get('topic','')}:  ")
        r_ut.bold = True
        p_u.add_run(f"추가 필요 정보 → {u.get('needed','')}")
        p_u.paragraph_format.left_indent = Inches(0.2)
    doc.add_paragraph()

    # ── F. 시사점 ────────────────────────────────────────────────────────────
    doc.add_paragraph("F. STRATEGIC IMPLICATIONS", style="JD_H2")
    impl = synthesis.get("implications", {})
    for label, key in [("경쟁사", "competitors"), ("파트너", "partners"), ("인재 시장", "talent")]:
        p_il = doc.add_paragraph(style="JD_H3")
        p_il.add_run(label)
        p_ib = doc.add_paragraph(style="JD_Body")
        p_ib.add_run(impl.get(key,""))
        p_ib.paragraph_format.left_indent = Inches(0.2)
        doc.add_paragraph()

    # 푸터 정보
    doc.add_page_break()
    p_foot = doc.add_paragraph(style="JD_Meta")
    p_foot.add_run(f"Generated by JD Analyst  ·  {date_str}  ·  {len(jds)} postings analyzed")
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ── UI ─────────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="hero">
  <div class="hero-eyebrow">Strategic Intelligence Tool</div>
  <div class="hero-title">Job Description Analyst</div>
  <div class="hero-desc">기업명을 입력하면 채용 페이지를 AI가 자동으로 탐색하고 전략 분석 리포트를 생성합니다.</div>
</div>
""", unsafe_allow_html=True)

with st.container():
    st.markdown('<div class="form-section">', unsafe_allow_html=True)
    _, col_c, _ = st.columns([1, 2, 1])
    with col_c:
        st.markdown('<span class="field-label">Company Name</span>', unsafe_allow_html=True)
        company = st.text_input("company", placeholder="Apple / Samsung / Google ...",
                                label_visibility="collapsed")
        st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

        st.markdown('<span class="field-label">JD URLs &nbsp;<span style="color:#2a2a2a;font-size:9px;letter-spacing:0.1em;">OPTIONAL — 분석할 채용공고 URL (여러 개면 줄바꿈으로 구분)</span></span>',
                    unsafe_allow_html=True)
        jd_urls_raw = st.text_area("jd_urls",
                                   placeholder="https://jobs.apple.com/en-us/details/...\nhttps://jobs.apple.com/en-us/details/...",
                                   height=100,
                                   label_visibility="collapsed")
        st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

        st.markdown('<span class="field-label">Filter &nbsp;<span style="color:#2a2a2a;font-size:9px;letter-spacing:0.1em;">OPTIONAL</span></span>',
                    unsafe_allow_html=True)
        filter_kw = st.text_input("filter",
                                  placeholder="Machine Learning / Marketing / Online Store ...",
                                  label_visibility="collapsed")
        st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)

    _, btn_c, _ = st.columns([1, 2, 1])
    with btn_c:
        run_btn = st.button("Start Analyzing", disabled=not company, use_container_width=True)

    st.markdown('</div>', unsafe_allow_html=True)

# ── 실행 ───────────────────────────────────────────────────────────────────────
if run_btn and company:
    genai.configure(api_key=GEMINI_API_KEY)

    log_placeholder = st.empty()
    log_lines = []

    def log(msg, style="active"):
        cls = {"ok":"log-ok","dim":"log-dim","active":"log-active"}.get(style,"log-active")
        log_lines.append(f'<div class="log-line {cls}">{msg}</div>')
        log_placeholder.markdown(
            f'<div class="log-box">{"".join(log_lines)}</div>',
            unsafe_allow_html=True)

    log(f"→ 분석 대상: {company}")
    log(f"→ 모델: Gemini 1.5 Pro")
    if filter_kw: log(f"→ 필터: {filter_kw}")

    # 사용자가 직접 입력한 JD URL 파싱
    manual_urls = [u.strip() for u in jd_urls_raw.strip().split("\n") if u.strip().startswith("http")]
    if manual_urls:
        log(f"→ 직접 입력 URL {len(manual_urls)}건 감지", "ok")
    log("─" * 36, "dim")

    try:
        # 0. URL 탐색
        careers_url = find_careers_url(company, log)
        st.markdown(f'<div class="url-found-box">채용 페이지 → {careers_url}</div>',
                    unsafe_allow_html=True)

        # 1. 공고 수집: 직접 입력 URL 우선, 없으면 자동 수집
        if manual_urls:
            log(f"◎ 직접 입력 URL {len(manual_urls)}건으로 분석 시작...")
            jobs = [{"title": f"JD {i+1}", "team": "", "location": "", "url": u}
                    for i, u in enumerate(manual_urls)]
            # URL에서 제목 추출 시도
            for job in jobs:
                slug = job["url"].rstrip("/").split("/")[-1].replace("-"," ").replace("_"," ")
                if slug and len(slug) > 3:
                    job["title"] = slug.title()
        else:
            jobs = collect_all_jobs(company, careers_url, filter_kw, log)

        # 2. 각 공고 심층 분석
        log(f"◎ {len(jobs)}건 공고 내부 심층 분석 시작...")
        jds = []
        for i, job in enumerate(jobs):
            jd = deep_analyze_jd(company, job, i+1, len(jobs), log)
            jds.append(jd)

        log(f"✓ 전체 {len(jds)}건 분석 완료", "ok")

        # 3. 전략 통합 분석
        synthesis = strategic_synthesis(company, filter_kw, jds, log)

        # 4. Word 문서 생성
        log("◎ Word 보고서 생성 중...")
        date_str  = datetime.now().strftime("%Y.%m.%d")
        docx_bytes = build_docx(company, filter_kw, careers_url, jds, synthesis, date_str)
        log("✓ 보고서 완성!", "ok")
        log("─" * 36, "dim")
        log("◎ 다운로드 준비 완료", "ok")

        st.markdown('<div class="result-section">', unsafe_allow_html=True)
        st.markdown(
            f'<div class="result-header">'
            f'<span>REPORT READY — {company.upper()}</span>'
            f'<span>{datetime.now().strftime("%Y.%m.%d %H:%M")}</span></div>',
            unsafe_allow_html=True)

        fname = f"{company.lower().replace(' ','_')}_jd_report_{datetime.now().strftime('%Y%m%d')}.docx"
        st.download_button(
            "↓  Word 리포트 다운로드",
            data=docx_bytes,
            file_name=fname,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    except Exception as e:
        err = str(e).lower()
        if "quota" in err or "429" in err or "resource_exhausted" in err:
            st.error("무료 사용량 한도 도달. 내일 자정(태평양 시간) 이후 다시 시도해주세요.")
        else:
            st.error(f"오류: {e}")

else:
    pass
